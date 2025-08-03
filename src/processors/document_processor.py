# src/processors/document_processor.py
import os
import json
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

# Document processing libraries
import docx
import PyPDF2
import openpyxl
from PIL import Image
import pytesseract
import cv2
import numpy as np
import pandas as pd
import email
from email import policy

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Multi-format document processor with OCR capabilities"""
    
    def __init__(self):
        self.supported_formats = {
            '.docx': self._process_docx,
            '.pdf': self._process_pdf,
            '.xlsx': self._process_xlsx,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.txt': self._process_txt,
            '.eml': self._process_eml,
            '.json': self._process_json,
            '.xml': self._process_xml,
            '.csv': self._process_csv
        }
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process a single file and extract all text content"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            logger.info(f"Processing {file_ext} file: {file_path}")
            
            result = self.supported_formats[file_ext](file_path)
            result.update({
                'file_path': file_path,
                'file_type': file_ext,
                'file_name': Path(file_path).name
            })
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return {
                'content': '',
                'error': str(e),
                'file_path': file_path,
                'file_type': file_ext if 'file_ext' in locals() else 'unknown'
            }
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX files including embedded images"""
        doc = docx.Document(file_path)
        
        # Extract text content
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract table content
        table_content = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Only add non-empty rows
                    table_content.append(' | '.join(row_data))
        
        # Extract embedded images and process with OCR
        image_text = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    # Extract and process embedded image
                    image_data = rel.target_part.blob
                    ocr_text = self._process_image_data(image_data)
                    if ocr_text:
                        image_text.append(ocr_text)
        except Exception as e:
            logger.warning(f"Could not extract embedded images: {str(e)}")
        
        return {
            'content': '\n'.join(text_content),
            'tables': table_content,
            'image_text': image_text,
            'total_paragraphs': len(text_content),
            'total_tables': len(doc.tables)
        }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files with text extraction and OCR fallback"""
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    # Try text extraction first
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        text_content.append(f"Page {page_num + 1}:\n{page_text}")
                    else:
                        # If no text found, use OCR on the page
                        logger.info(f"No text found on page {page_num + 1}, attempting OCR")
                        # Note: For production, you'd convert PDF page to image first
                        text_content.append(f"Page {page_num + 1}: [OCR processing needed]")
        
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return {'content': '', 'error': str(e)}
        
        return {
            'content': '\n\n'.join(text_content),
            'total_pages': len(pdf_reader.pages) if 'pdf_reader' in locals() else 0
        }
    
    def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Process Excel files including all sheets and embedded content"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_content = {}
        all_text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    sheet_data.append(row_text)
            
            sheets_content[sheet_name] = sheet_data
            all_text.extend(sheet_data)
        
        return {
            'content': '\n'.join(all_text),
            'sheets': sheets_content,
            'sheet_names': list(workbook.sheetnames),
            'total_sheets': len(workbook.sheetnames)
        }
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Enhanced image processing with multiple OCR attempts"""
        try:
            # Load and preprocess image
            image = cv2.imread(file_path)
            if image is None:
                raise ValueError("Could not load image")
            
            # Try multiple OCR configurations for better accuracy
            ocr_configs = [
                '--oem 3 --psm 6',  # Default
                '--oem 3 --psm 8',  # Single word
                '--oem 3 --psm 7',  # Single text line
                '--oem 3 --psm 11', # Sparse text
                '--oem 3 --psm 13'  # Raw line
            ]
            
            best_text = ""
            best_confidence = 0
            
            for config in ocr_configs:
                try:
                    processed_image = self._preprocess_image(image)
                    
                    # Get OCR result with confidence
                    ocr_data = pytesseract.image_to_data(processed_image, config=config, output_type=pytesseract.Output.DICT)
                    
                    # Calculate average confidence
                    confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                    
                    # Get text
                    text = pytesseract.image_to_string(processed_image, config=config).strip()
                    
                    # Keep best result
                    if avg_confidence > best_confidence and len(text) > len(best_text):
                        best_text = text
                        best_confidence = avg_confidence
                        
                except Exception as e:
                    logger.warning(f"OCR config {config} failed: {str(e)}")
                    continue
            
            # Clean up OCR text
            cleaned_text = self._clean_ocr_text(best_text)
            
            return {
                'content': cleaned_text,
                'image_size': image.shape[:2],
                'ocr_confidence': best_confidence,
                'ocr_method': 'multi-config'
            }
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            return {'content': '', 'error': str(e)}
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean up OCR text for better accuracy"""
        if not text:
            return ""
        
        # Common OCR corrections
        corrections = {
            'gate': 'date',
            'Beneticiary': 'Beneficiary',
            'Bene:iciary': 'Beneficiary',
            'Bene ficiary': 'Beneficiary',
            'Arnount': 'Amount',
            'Am0unt': 'Amount',
            'Va|ue': 'Value',
            'V4lue': 'Value'
        }
        
        cleaned = text
        for wrong, correct in corrections.items():
            cleaned = cleaned.replace(wrong, correct)
        
        # Remove excessive whitespace
        cleaned = ' '.join(cleaned.split())
        
        return cleaned
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'content': content,
                'line_count': len(content.split('\n'))
            }
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return {
                        'content': content,
                        'encoding_used': encoding,
                        'line_count': len(content.split('\n'))
                    }
                except UnicodeDecodeError:
                    continue
            
            return {'content': '', 'error': 'Could not decode file'}
    
    def _process_eml(self, file_path: str) -> Dict[str, Any]:
        """Process email files"""
        try:
            with open(file_path, 'rb') as file:
                msg = email.message_from_bytes(file.read(), policy=policy.default)
            
            # Extract email components
            subject = msg.get('Subject', '')
            sender = msg.get('From', '')
            recipient = msg.get('To', '')
            date = msg.get('Date', '')
            
            # Extract body
            body = ''
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_content()
            else:
                body = msg.get_content()
            
            content = f"Subject: {subject}\nFrom: {sender}\nTo: {recipient}\nDate: {date}\n\nBody:\n{body}"
            
            return {
                'content': content,
                'subject': subject,
                'sender': sender,
                'recipient': recipient
            }
            
        except Exception as e:
            logger.error(f"Error processing email: {str(e)}")
            return {'content': '', 'error': str(e)}
    
    def _process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            return {
                'content': content,
                'json_data': data,
                'keys': list(data.keys()) if isinstance(data, dict) else None
            }
            
        except Exception as e:
            logger.error(f"Error processing JSON: {str(e)}")
            return {'content': '', 'error': str(e)}
    
    def _process_xml(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract text content from XML
            content = ET.tostring(root, encoding='unicode', method='text')
            
            return {
                'content': content,
                'root_tag': root.tag,
                'element_count': len(list(root.iter()))
            }
            
        except Exception as e:
            logger.error(f"Error processing XML: {str(e)}")
            return {'content': '', 'error': str(e)}
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = df.to_string(index=False)
            
            return {
                'content': content,
                'columns': list(df.columns),
                'row_count': len(df),
                'column_count': len(df.columns)
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV: {str(e)}")
            return {'content': '', 'error': str(e)}
    
    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Enhanced image preprocessing for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Resize image if too small (improves OCR accuracy)
        height, width = gray.shape
        if height < 300 or width < 300:
            scale_factor = max(300/height, 300/width)
            new_height = int(height * scale_factor)
            new_width = int(width * scale_factor)
            gray = cv2.resize(gray, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
        
        # Apply Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (3, 3), 0)
        
        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(blurred)
        
        # Apply morphological operations to clean up
        kernel = np.ones((2, 2), np.uint8)
        cleaned = cv2.morphologyEx(enhanced, cv2.MORPH_CLOSE, kernel)
        
        # Apply threshold with Otsu's method
        _, thresh = cv2.threshold(cleaned, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return thresh
    
    def _get_ocr_confidence(self, image: np.ndarray) -> float:
        """Get OCR confidence score"""
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            return sum(confidences) / len(confidences) if confidences else 0
        except:
            return 0
    
    def _process_image_data(self, image_data: bytes) -> str:
        """Process image data from embedded content"""
        try:
            # Convert bytes to image
            image_array = np.frombuffer(image_data, np.uint8)
            image = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
            
            if image is not None:
                processed_image = self._preprocess_image(image)
                ocr_text = pytesseract.image_to_string(processed_image, config='--oem 3 --psm 6')
                return ocr_text.strip()
            
        except Exception as e:
            logger.warning(f"Could not process embedded image: {str(e)}")
        
        return ""

# Usage example
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample file
    test_file = "sample.docx"  # Replace with actual file path
    result = processor.process_file(test_file)
    
    print(f"Extracted content: {result['content'][:500]}...")
    print(f"File type: {result['file_type']}")